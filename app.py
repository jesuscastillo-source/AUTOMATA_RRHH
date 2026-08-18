# -*- coding: utf-8 -*-
"""
Generador de Documentos RRHH
=============================
App web (Streamlit) que reemplaza el notebook de Colab original.
5 herramientas, cada una con subida manual de Excel de datos + plantilla(s).

Cómo correr localmente:
    pip install -r requirements.txt
    streamlit run app.py
"""

import io
import os
import re
import zipfile
import datetime
from datetime import timedelta

import pandas as pd
import streamlit as st
from openpyxl import load_workbook
from openpyxl.utils.datetime import from_excel
from docx import Document
from docx.shared import Pt
from num2words import num2words
import holidays


# =========================================================
# HELPERS GENERALES
# =========================================================

def format_fecha(valor):
    """Formatea cualquier valor de fecha (date, datetime, número serial de Excel
    o string en varios formatos) a DD/MM/YYYY."""
    if valor in (None, ""):
        return ""
    if isinstance(valor, (datetime.datetime, datetime.date)):
        return valor.strftime("%d/%m/%Y")
    if isinstance(valor, (int, float)):
        try:
            return from_excel(valor).strftime("%d/%m/%Y")
        except Exception:
            pass
    if isinstance(valor, str):
        valor = valor.strip()
        for fmt in ("%d-%m-%Y", "%d/%m/%Y", "%Y-%m-%d"):
            try:
                return datetime.datetime.strptime(valor, fmt).strftime("%d/%m/%Y")
            except Exception:
                continue
        try:
            dt = pd.to_datetime(valor, dayfirst=True, errors="coerce")
            if not pd.isnull(dt):
                return dt.strftime("%d/%m/%Y")
        except Exception:
            pass
    return str(valor)


def monto_a_texto(valor, prefijo=""):
    """1234567 -> '1.234.567 (UN MILLON ... PESOS)' (con prefijo opcional, ej '$')."""
    valor_limpio = str(valor).replace(".", "").replace(",", "")
    valor_int = int(float(valor_limpio))
    valor_numero = prefijo + "{:,}".format(valor_int).replace(",", ".")
    valor_letras = num2words(valor_int, lang="es").upper() + " PESOS"
    return f"{valor_numero} ({valor_letras})"


def sueldo_en_letras(valor):
    """Igual que monto_a_texto pero retorna (numero, letras) por separado."""
    valor_int = int(float(str(valor).replace(".", "").replace(",", "")))
    numero = "{:,}".format(valor_int).replace(",", ".")
    letras = num2words(valor_int, lang="es").upper() + " PESOS"
    return numero, letras


def crear_zip(archivos: dict) -> bytes:
    """archivos: {nombre_archivo: bytes}. Soporta rutas con '/' para subcarpetas."""
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for nombre, contenido in archivos.items():
            zf.writestr(nombre, contenido)
    buffer.seek(0)
    return buffer.getvalue()


def guardar_docx_bytes(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def guardar_xlsx_bytes(wb) -> bytes:
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# =========================================================
# LECTURA DE EXCEL (variantes según lo que necesitaba cada tool)
# =========================================================

def leer_excel_formato(file_bytes, date_headers, money_headers=None, money_prefijo=""):
    """Lee Excel preservando negrita/tamaño de fuente por celda (para mergefields
    que necesitan mantener formato). Devuelve lista de dicts: header -> (valor, bold, size)."""
    money_headers = money_headers or []
    wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
    ws = wb.active
    excel_data = []

    for row in ws.iter_rows(min_row=2, values_only=False):
        if all(c.value in (None, "") for c in row):
            continue
        row_data = {}
        for cell in row:
            header_raw = ws.cell(row=1, column=cell.column).value
            if header_raw is None:
                continue
            header = str(header_raw).strip().upper()
            valor = cell.value
            bold = cell.font.bold if cell.font.bold is not None else False
            size = cell.font.sz if cell.font.sz is not None else None

            if header in date_headers:
                valor = format_fecha(valor)
            elif header in money_headers and valor not in (None, ""):
                try:
                    valor = monto_a_texto(valor, prefijo=money_prefijo)
                except Exception:
                    valor = str(valor)

            row_data[header] = (valor, bold, size)

        nombre = row_data.get("NOMBRE COMPLETO", ("", False, None))[0]
        if nombre is not None and str(nombre).strip() != "":
            excel_data.append(row_data)

    return excel_data


def leer_excel_simple(file_bytes, date_headers):
    """Lee Excel a strings planos (sin negrita/tamaño) — para anexos de continuidad."""
    wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
    ws = wb.active
    excel_data = []

    for row in ws.iter_rows(min_row=2, values_only=False):
        if all(c.value in (None, "") for c in row):
            continue
        row_data = {}
        for cell in row:
            header_raw = ws.cell(row=1, column=cell.column).value
            if header_raw is None:
                continue
            header = str(header_raw).strip().upper()
            valor = cell.value
            if header in date_headers and valor is not None:
                valor = format_fecha(valor)
            row_data[header] = "" if valor is None else str(valor)

        if row_data.get("NOMBRE COMPLETO", "").strip():
            excel_data.append(row_data)

    return excel_data


def leer_excel_obrero_capataz(file_bytes, date_headers):
    """Variante especial: separa SUELDO CAPATAZ en número y letras."""
    wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
    ws = wb.active
    headers = [str(c.value).strip().upper() if c.value is not None else "" for c in ws[1]]
    excel_data = []

    for row in ws.iter_rows(min_row=2):
        if all(c.value in (None, "") for c in row):
            continue
        row_data = {}
        for header, cell in zip(headers, row):
            if not header:
                continue
            valor = cell.value

            if header in date_headers:
                row_data[header] = format_fecha(valor)
                continue

            if header == "SUELDO CAPATAZ" and valor not in (None, ""):
                try:
                    numero, letras = sueldo_en_letras(valor)
                    row_data["SUELDO CAPATAZ"] = numero
                    row_data["SUELDO CAPATAZ LETRAS"] = letras
                except Exception:
                    row_data["SUELDO CAPATAZ"] = str(valor)
                    row_data["SUELDO CAPATAZ LETRAS"] = ""
                continue

            row_data[header] = "" if valor is None else str(valor)

        if row_data.get("NOMBRE COMPLETO", "").strip():
            excel_data.append(row_data)

    return excel_data


# =========================================================
# REEMPLAZO DE MERGEFIELDS EN WORD
# =========================================================

MERGEFIELD_RE = re.compile(r"«([^»]*)»")


def replace_mergefield_con_formato(doc, replacements):
    """replacements: header (en MAYÚSCULAS) -> (valor, bold, size).
    Busca cualquier «campo» sin importar si el molde lo tiene en mayúsculas,
    minúsculas o mezclado — siempre lo compara en mayúsculas."""
    def procesar(parrafo):
        for run in parrafo.runs:
            estado = {}

            def sub(m):
                campo = m.group(1).strip().upper()
                if campo in replacements:
                    valor, bold, size = replacements[campo]
                    estado["bold"] = bold
                    estado["size"] = size
                    return str(valor)
                return m.group(0)

            nuevo_texto = MERGEFIELD_RE.sub(sub, run.text)
            if nuevo_texto != run.text:
                run.text = nuevo_texto
                if "bold" in estado:
                    run.bold = estado["bold"]
                if estado.get("size"):
                    run.font.size = Pt(estado["size"])

    for p in doc.paragraphs:
        procesar(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    procesar(p)


def replace_mergefields_simple(doc, replacements, bold_keys=None):
    """replacements: header (en MAYÚSCULAS) -> valor (string plano).
    Misma lógica insensible a mayúsculas/minúsculas que la versión con formato."""
    bold_keys = bold_keys or set()

    def procesar(parrafo):
        for run in parrafo.runs:
            estado = {}

            def sub(m):
                campo = m.group(1).strip().upper()
                if campo in replacements:
                    if campo in bold_keys:
                        estado["bold"] = True
                    return str(replacements[campo])
                return m.group(0)

            nuevo_texto = MERGEFIELD_RE.sub(sub, run.text)
            if nuevo_texto != run.text:
                run.text = nuevo_texto
                if estado.get("bold"):
                    run.bold = True

    for p in doc.paragraphs:
        procesar(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    procesar(p)


# =========================================================
# CÁLCULO DE FINIQUITOS (feriado proporcional + días inhábiles)
# =========================================================

DIAS_SEMANA_ES = {
    "Monday": "lunes", "Tuesday": "martes", "Wednesday": "miércoles",
    "Thursday": "jueves", "Friday": "viernes", "Saturday": "sábado", "Sunday": "domingo",
}


def calcular_feriado_proporcional(inicio, fin):
    dias_trabajados = (fin - inicio).days + 1
    if dias_trabajados < 30:
        return None, dias_trabajados, 0, 0, 0
    anios = dias_trabajados // 365
    resto = dias_trabajados % 365
    meses = resto // 30
    dias = resto % 30
    total = anios * 15 + meses * 1.25 + dias * (1.25 / 30)
    return round(total, 1), dias_trabajados, anios, meses, dias


def calcular_dias_inhabiles(fecha_fin, dias_proporcionales):
    parte_entera = int(dias_proporcionales)
    fraccion = dias_proporcionales - parte_entera
    dias_a_contar = parte_entera + (1 if fraccion > 0 else 0)

    fecha_cursor = fecha_fin
    dias_habiles = 0
    dias_inhabiles = 0
    calendario = []

    years = [fecha_fin.year, fecha_fin.year + 1]
    feriados = holidays.Chile(years=years)

    while dias_habiles < dias_a_contar:
        fecha_cursor += timedelta(days=1)
        es_feriado = fecha_cursor in feriados
        es_fin_semana = fecha_cursor.weekday() >= 5
        if es_fin_semana or es_feriado:
            dias_inhabiles += 1
            calendario.append((fecha_cursor, "inhábil"))
        else:
            dias_habiles += 1
            calendario.append((fecha_cursor, "hábil"))

    return dias_inhabiles, calendario


# =========================================================
# UI STREAMLIT
# =========================================================

st.set_page_config(page_title="Automatización RR.HH.", page_icon="🤖", layout="wide")

# --- Banner ---
import base64

def _banner_base64():
    with open(os.path.join(os.path.dirname(__file__), "assets", "banner.jpg"), "rb") as f:
        return base64.b64encode(f.read()).decode()

st.markdown(
    """
    <style>
    .block-container { padding-top: 1.5rem; max-width: 1180px; }

    .hero-banner {
        border-radius: 20px;
        overflow: hidden;
        margin-bottom: 1.8rem;
        box-shadow: 0 8px 24px rgba(0,0,0,0.35);
    }
    .hero-banner img { width: 100%; height: auto; display: block; }

    /* --- Pestañas como botones tipo "pill" --- */
    .stTabs [role="tablist"] {
        gap: 10px;
        flex-wrap: wrap;
        border-bottom: none !important;
        box-shadow: none !important;
        padding-bottom: 0.6rem;
    }
    .stTabs [data-testid="stTab"] .react-aria-SelectionIndicator {
        display: none !important;
    }
    .stTabs [data-testid="stTab"]::before,
    .stTabs [data-testid="stTab"]::after {
        display: none !important;
        content: none !important;
    }
    .stTabs [data-testid="stTab"] {
        height: auto;
        background-color: rgba(255,255,255,0.05) !important;
        border: 1px solid #2A3742 !important;
        border-radius: 999px !important;
        padding: 0.6rem 1.25rem !important;
        margin: 0 !important;
        box-shadow: none !important;
        outline: none !important;
        transition: all 0.15s ease;
    }
    .stTabs [data-testid="stTab"] p {
        font-weight: 600 !important;
        font-size: 0.92rem !important;
        color: #9AA5AE;
        margin: 0 !important;
    }
    .stTabs [data-testid="stTab"]:hover {
        background-color: rgba(127,184,143,0.12) !important;
        border-color: #5F9E72 !important;
        box-shadow: none !important;
    }
    .stTabs [data-testid="stTab"]:hover p { color: #DCE3E8; }
    .stTabs [data-testid="stTab"][aria-selected="true"] {
        background: linear-gradient(135deg, #8FCB9E 0%, #5F9E72 100%) !important;
        border-color: #7FB88F !important;
        box-shadow: 0 4px 14px rgba(127,184,143,0.30) !important;
    }
    .stTabs [data-testid="stTab"][aria-selected="true"]:focus,
    .stTabs [data-testid="stTab"][aria-selected="true"]:focus-visible {
        box-shadow: 0 4px 14px rgba(127,184,143,0.30) !important;
        outline: none !important;
    }
    .stTabs [data-testid="stTab"][aria-selected="true"] p { color: #0F151C !important; }

    .stButton>button {
        border-radius: 10px;
        font-weight: 600;
        border: 1px solid #5F9E72;
    }
    .stButton>button:hover { border-color: #7FB88F; color: #7FB88F; }
    </style>
    """,
    unsafe_allow_html=True,
)

st.markdown(
    f'<div class="hero-banner"><img src="data:image/jpeg;base64,{_banner_base64()}" alt="Automatización de Documentación RR.HH."></div>',
    unsafe_allow_html=True,
)

tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "📑 Contratos",
    "🧮 Cálculo Finiquitos",
    "📋 Finiquito + Decl. Jurada",
    "➕ Anexo Continuidad",
    "⬆️ Anexo Obrero→Capataz",
])


# ---------------------------------------------------------
# TAB 1: CONTRATOS
# ---------------------------------------------------------
with tab1:
    st.subheader("Generador de Contratos")
    with st.expander("Columnas esperadas en el Excel"):
        st.markdown(
            "- `NOMBRE COMPLETO`\n"
            "- `FECHA ACTUAL`, `FECHA DE NACIMIENTO`, `INICIO CONTRARO`, `TÉRMINO CONTRARO` (fechas)\n"
            "- `SUELDO BASE` (número)\n"
            "- Cualquier otra columna se reemplaza tal cual en el molde como `«NOMBRE_COLUMNA»`"
        )

    c1, c2 = st.columns(2)
    excel_c = c1.file_uploader("Excel de datos", type=["xlsx"], key="c_excel")
    molde_c = c2.file_uploader("Plantilla Word (molde de contrato)", type=["docx"], key="c_molde")

    if st.button("Generar contratos", key="btn_c"):
        if not excel_c or not molde_c:
            st.error("Sube el Excel y la plantilla primero.")
        else:
            try:
                date_headers = ["FECHA ACTUAL", "FECHA DE NACIMIENTO", "INICIO CONTRARO", "TÉRMINO CONTRARO"]
                excel_data = leer_excel_formato(excel_c.getvalue(), date_headers, money_headers=["SUELDO BASE"])

                if not excel_data:
                    st.warning("No se encontraron filas válidas (revisa que exista la columna NOMBRE COMPLETO).")
                else:
                    archivos = {}
                    molde_bytes = molde_c.getvalue()
                    for idx, row_data in enumerate(excel_data):
                        doc = Document(io.BytesIO(molde_bytes))
                        replace_mergefield_con_formato(doc, row_data)
                        numero = str(idx + 1).zfill(3)
                        nombre = str(row_data["NOMBRE COMPLETO"][0]).replace(" ", "_")
                        archivos[f"{numero}_CTO_{nombre}.docx"] = guardar_docx_bytes(doc)

                    zip_bytes = crear_zip(archivos)
                    st.success(f"{len(archivos)} contrato(s) generado(s).")
                    st.download_button("⬇️ Descargar ZIP", zip_bytes, "Contratos_Generados.zip", "application/zip")
            except Exception as e:
                st.error(f"Error generando los contratos: {e}")


# ---------------------------------------------------------
# TAB 2: CÁLCULO DE FINIQUITOS (planillas)
# ---------------------------------------------------------
with tab2:
    st.subheader("Cálculo de Finiquitos (planilla Excel)")
    with st.expander("Columnas esperadas en el Excel de datos"):
        st.markdown(
            "- `NOMBRE COMPLETO`\n"
            "- `CARNÉ DE IDENTIDAD N°`\n"
            "- `INICIO CONTRARO`, `FIN CONTRAT` (fechas)\n"
            "- `SUELDO BASE`\n"
            "- `CAUSAL TÉRMINO`\n\n"
            "La plantilla Excel debe tener las celdas: `D6` nombre, `D7` RUT, `D8` PEE, "
            "`D14` causal, `D17` inicio, `D18` fin, `L20` días inhábiles, `D39` sueldo base."
        )

    c1, c2 = st.columns(2)
    excel_f = c1.file_uploader("Excel de datos", type=["xlsx"], key="f_excel")
    plantilla_f = c2.file_uploader("Plantilla Excel (cálculo finiquito)", type=["xlsx"], key="f_plantilla")

    if st.button("Calcular y generar planillas", key="btn_f"):
        if not excel_f or not plantilla_f:
            st.error("Sube el Excel de datos y la plantilla primero.")
        else:
            try:
                df = pd.read_excel(io.BytesIO(excel_f.getvalue()))
                df.columns = df.columns.str.strip().str.upper()

                requeridas = ["NOMBRE COMPLETO", "CARNÉ DE IDENTIDAD N°", "INICIO CONTRARO",
                              "FIN CONTRAT", "SUELDO BASE", "CAUSAL TÉRMINO"]
                faltantes = [c for c in requeridas if c not in df.columns]
                if faltantes:
                    st.error(f"Faltan columnas en el Excel: {', '.join(faltantes)}")
                else:
                    plantilla_bytes = plantilla_f.getvalue()
                    archivos = {}
                    resumen = []

                    for index, row in df.iterrows():
                        nombre = row["NOMBRE COMPLETO"]
                        rut = row["CARNÉ DE IDENTIDAD N°"]
                        fecha_inicio = pd.to_datetime(row["INICIO CONTRARO"])
                        fecha_fin = pd.to_datetime(row["FIN CONTRAT"])
                        sueldo_base = row["SUELDO BASE"]
                        causal = row["CAUSAL TÉRMINO"]

                        dias_prop, dias_trab, anios, meses, dias = calcular_feriado_proporcional(fecha_inicio, fecha_fin)

                        if dias_prop is None:
                            dias_inhabiles = 0
                        else:
                            dias_inhabiles, _calendario = calcular_dias_inhabiles(fecha_fin, dias_prop)

                        resumen.append({
                            "Nombre": nombre, "Días trabajados": dias_trab,
                            "Días proporcionales": dias_prop, "Días inhábiles": dias_inhabiles,
                        })

                        wb = load_workbook(io.BytesIO(plantilla_bytes))
                        wb.calculation.fullCalcOnLoad = True
                        ws = wb.active
                        ws["D6"] = nombre
                        ws["D7"] = rut
                        ws["D8"] = "PEE"
                        ws["D14"] = causal
                        ws["D17"] = fecha_inicio.strftime("%d-%m-%Y")
                        ws["D18"] = fecha_fin.strftime("%d-%m-%Y")
                        ws["L20"] = dias_inhabiles
                        ws["D39"] = sueldo_base

                        numero = str(index + 1).zfill(2)
                        nombre_archivo = f"{numero}_CALCULO_FINIQUITO_{str(nombre).replace(' ', '_')}.xlsx"
                        archivos[nombre_archivo] = guardar_xlsx_bytes(wb)

                    zip_bytes = crear_zip(archivos)
                    st.success(f"{len(archivos)} planilla(s) generada(s).")
                    st.dataframe(pd.DataFrame(resumen), use_container_width=True)
                    st.download_button("⬇️ Descargar ZIP", zip_bytes, "CALCULO_FINIQU_GENERADOS.zip", "application/zip")
            except Exception as e:
                st.error(f"Error calculando los finiquitos: {e}")


# ---------------------------------------------------------
# TAB 3: FINIQUITO + DECLARACIÓN JURADA
# ---------------------------------------------------------
with tab3:
    st.subheader("Finiquito + Declaración Jurada")
    with st.expander("Columnas esperadas en el Excel"):
        st.markdown(
            "- `NOMBRE COMPLETO`, `CARNÉ DE IDENTIDAD N°`\n"
            "- `FECHA ACTUAL`, `INICIO CONTRATO` / `INICIO CONTRARO`, `FIN CONTRAT` / `FIN CONTRATO` (fechas)\n"
            "- `MONTO FINIQUITO` (número)"
        )

    c1, c2, c3 = st.columns(3)
    excel_j = c1.file_uploader("Excel de datos", type=["xlsx"], key="j_excel")
    molde_finiquito = c2.file_uploader("Plantilla Finiquito (Word)", type=["docx"], key="j_finiquito")
    molde_jurada = c3.file_uploader("Plantilla Declaración Jurada (Word)", type=["docx"], key="j_jurada")

    if st.button("Generar finiquitos + declaraciones", key="btn_j"):
        if not excel_j or not molde_finiquito or not molde_jurada:
            st.error("Sube el Excel y las dos plantillas primero.")
        else:
            try:
                date_headers = ["FECHA ACTUAL", "INICIO CONTRATO", "INICIO CONTRARO", "FIN CONTRAT", "FIN CONTRATO"]
                excel_data = leer_excel_formato(
                    excel_j.getvalue(), date_headers,
                    money_headers=["MONTO FINIQUITO"], money_prefijo="$"
                )

                if not excel_data:
                    st.warning("No se encontraron filas válidas.")
                else:
                    archivos = {}
                    finiquito_bytes = molde_finiquito.getvalue()
                    jurada_bytes = molde_jurada.getvalue()
                    contador = 1

                    for row_data in excel_data:
                        doc = Document(io.BytesIO(finiquito_bytes))
                        replace_mergefield_con_formato(doc, row_data)
                        nombre_trabajador = str(row_data["NOMBRE COMPLETO"][0]).replace(" ", "_")
                        archivos[f"Finiquitos/{contador:03d}_Finiquito_{nombre_trabajador}.docx"] = guardar_docx_bytes(doc)
                        contador += 1

                    contador = 1
                    for row_data in excel_data:
                        jurada_data = {
                            "NOMBRE COMPLETO": row_data["NOMBRE COMPLETO"],
                            "CARNÉ DE IDENTIDAD N°": row_data.get("CARNÉ DE IDENTIDAD N°", ("", False, None)),
                        }
                        doc = Document(io.BytesIO(jurada_bytes))
                        replace_mergefield_con_formato(doc, jurada_data)
                        nombre_trabajador = str(row_data["NOMBRE COMPLETO"][0]).replace(" ", "_")
                        archivos[f"Declaraciones_Juradas/{contador:03d}_Decl_Jurada_{nombre_trabajador}.docx"] = guardar_docx_bytes(doc)
                        contador += 1

                    zip_bytes = crear_zip(archivos)
                    st.success(f"{len(excel_data)} finiquito(s) + declaración(es) generados.")
                    st.download_button("⬇️ Descargar ZIP", zip_bytes, "Finiquitos_Generados.zip", "application/zip")
            except Exception as e:
                st.error(f"Error generando finiquitos/declaraciones: {e}")


# ---------------------------------------------------------
# TAB 4: ANEXO CONTINUIDAD
# ---------------------------------------------------------
with tab4:
    st.subheader("Anexo de Continuidad")
    with st.expander("Columnas esperadas en el Excel"):
        st.markdown(
            "- `NOMBRE COMPLETO`\n"
            "- `FECHA ACTUAL`, `FECHA INICIO CONTRATO`, `FECHA FIN CONTRATO ANEXO` (fechas)"
        )

    c1, c2 = st.columns(2)
    excel_a = c1.file_uploader("Excel de datos", type=["xlsx"], key="a_excel")
    molde_a = c2.file_uploader("Plantilla Word (anexo)", type=["docx"], key="a_molde")

    if st.button("Generar anexos", key="btn_a"):
        if not excel_a or not molde_a:
            st.error("Sube el Excel y la plantilla primero.")
        else:
            try:
                date_headers = ["FECHA ACTUAL", "FECHA INICIO CONTRATO", "FECHA FIN CONTRATO ANEXO"]
                excel_data = leer_excel_simple(excel_a.getvalue(), date_headers)

                if not excel_data:
                    st.warning("No se encontraron filas válidas.")
                else:
                    archivos = {}
                    molde_bytes = molde_a.getvalue()
                    for idx, row_data in enumerate(excel_data, start=1):
                        doc = Document(io.BytesIO(molde_bytes))
                        replace_mergefields_simple(doc, row_data)
                        nombre = row_data["NOMBRE COMPLETO"].replace(" ", "_")
                        archivos[f"{idx:02d}_{nombre}_ANEXO.docx"] = guardar_docx_bytes(doc)

                    zip_bytes = crear_zip(archivos)
                    st.success(f"{len(archivos)} anexo(s) generado(s).")
                    st.download_button("⬇️ Descargar ZIP", zip_bytes, "Anexos_Contrato.zip", "application/zip")
            except Exception as e:
                st.error(f"Error generando los anexos: {e}")


# ---------------------------------------------------------
# TAB 5: ANEXO OBRERO -> CAPATAZ
# ---------------------------------------------------------
with tab5:
    st.subheader("Anexo Obrero → Capataz")
    with st.expander("Columnas esperadas en el Excel"):
        st.markdown(
            "- `NOMBRE COMPLETO`\n"
            "- `FECHA ACTUAL`, `FECHA INICIO CONTRATO`, `FECHA FIN CONTRATO ANEXO`, `FECHA OB CPT DESDE` (fechas)\n"
            "- `SUELDO CAPATAZ` (número; se genera automáticamente `«SUELDO CAPATAZ»` y `«SUELDO CAPATAZ LETRAS»`)"
        )

    c1, c2 = st.columns(2)
    excel_o = c1.file_uploader("Excel de datos", type=["xlsx"], key="o_excel")
    molde_o = c2.file_uploader("Plantilla Word (anexo obrero→capataz)", type=["docx"], key="o_molde")

    if st.button("Generar modificaciones de contrato", key="btn_o"):
        if not excel_o or not molde_o:
            st.error("Sube el Excel y la plantilla primero.")
        else:
            try:
                date_headers = ["FECHA ACTUAL", "FECHA INICIO CONTRATO", "FECHA FIN CONTRATO ANEXO", "FECHA OB CPT DESDE"]
                excel_data = leer_excel_obrero_capataz(excel_o.getvalue(), date_headers)

                if not excel_data:
                    st.warning("No se encontraron filas válidas.")
                else:
                    archivos = {}
                    molde_bytes = molde_o.getvalue()
                    for i, row_data in enumerate(excel_data, start=1):
                        doc = Document(io.BytesIO(molde_bytes))
                        replace_mergefields_simple(doc, row_data, bold_keys={"SUELDO CAPATAZ LETRAS"})
                        nombre = row_data["NOMBRE COMPLETO"].replace(" ", "_")
                        archivos[f"{i:03d}_MOD_CONTRATO_{nombre}.docx"] = guardar_docx_bytes(doc)

                    zip_bytes = crear_zip(archivos)
                    st.success(f"{len(archivos)} documento(s) generado(s).")
                    st.download_button("⬇️ Descargar ZIP", zip_bytes, "Modificacion_Contrato_OB_CPT.zip", "application/zip")
            except Exception as e:
                st.error(f"Error generando las modificaciones: {e}")
